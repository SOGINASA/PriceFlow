"""Извлечение из DOCX (ТЗ 4.2). python-docx.

ВАЖНО: принять все tracked changes — работать с финальной версией текста.

python-docx читает только run'ы — прямые потомки абзаца (`<w:p>/<w:r>`), поэтому
вставки `<w:ins><w:r>...` НЕ попадают в `paragraph.text`/`cell.text` и теряются.
Решение: текст собираем напрямую по узлам `<w:t>`. Это автоматически даёт
«принять все изменения»: вставки (`<w:ins>` → содержит `<w:t>`) включаются, а
удаления (`<w:del>` → текст в `<w:delText>`, не `<w:t>`) исключаются. Перемещения
(`<w:moveFrom>`) вырезаем заранее, иначе перемещённый текст задвоится.
"""
import logging

from services.extractors import ExtractResult, RawRow
from services.extractors.row_parser import parse_table_row, looks_like_header, parse_text_line

logger = logging.getLogger(__name__)


def _accept_tracked_changes(document):
    """Принять отслеживаемые изменения (ТЗ 4.2): вырезать отклонённый контент.

    Удаляем поддеревья `<w:del>` и `<w:moveFrom>`, но только содержательные
    (с run'ами) — пустые маркеры свойств (`<w:trPr><w:del/>` и т.п.) не трогаем,
    чтобы не сломать структуру. Вставки остаются и попадут в текст через `<w:t>`.
    """
    try:
        from docx.oxml.ns import qn
        body = document.element.body
        for tag in (qn('w:del'), qn('w:moveFrom')):
            for el in list(body.iter(tag)):
                if el.find(qn('w:r')) is None:   # маркер свойства, не контент
                    continue
                parent = el.getparent()
                if parent is not None:
                    parent.remove(el)
    except Exception as e:  # noqa: BLE001
        logger.warning('accept tracked changes: %s', e)


def _xml_text(element) -> str:
    """Собрать видимый текст по узлам `<w:t>` (включая вложенные в `<w:ins>`)."""
    from docx.oxml.ns import qn
    return ''.join(t.text or '' for t in element.iter(qn('w:t')))


def _para_text(paragraph) -> str:
    """Текст абзаца с учётом принятых вставок (в обход paragraph.text)."""
    return _xml_text(paragraph._p).strip()


def _cell_text(cell) -> str:
    """Текст ячейки с учётом принятых вставок; абзацы внутри ячейки — через пробел."""
    return ' '.join(_xml_text(p._p) for p in cell.paragraphs).strip()


def extract(file_path: str) -> ExtractResult:
    result = ExtractResult()
    try:
        from docx import Document
    except ImportError:
        result.warnings.append('python-docx не установлен')
        return result

    try:
        document = Document(file_path)
        _accept_tracked_changes(document)

        # таблицы
        for table in document.tables:
            for row in table.rows:
                cells = [_cell_text(c) for c in row.cells]
                cells = [c for c in cells if c]
                if not cells or looks_like_header(cells):
                    continue
                parsed = parse_table_row(cells)
                if parsed:
                    result.rows.append(parsed)

        # абзацы (на случай прайса без таблиц)
        for p in document.paragraphs:
            text = _para_text(p)
            if not text:
                continue
            result.raw_text += text + '\n'
            if not document.tables:
                parsed = parse_text_line(text)
                if parsed:
                    result.rows.append(parsed)
    except Exception as e:  # noqa: BLE001
        logger.exception('docx extract failed')
        result.warnings.append(f'Ошибка чтения DOCX: {e}')
    return result
